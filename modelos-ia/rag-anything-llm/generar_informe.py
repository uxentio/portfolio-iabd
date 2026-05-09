"""
Script para generar Informe_MauiRAG.docx con formato profesional.
Usa python-docx para crear el documento con portada, indice, headings,
capturas incrustadas y bloques de codigo.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = r"D:\Tarea 5 RAG"
CAPTURAS_DIR = os.path.join(BASE_DIR, "Entrega", "Capturas")

# Mapeo de marcas [CAPTURA: ...] a archivos de imagen
CAPTURA_MAP = {
    "pantalla Config con los campos": "config.png",
    "pantalla Workspaces": "workspaces.png",
    "pantalla Documentos con un PDF seleccionado y el checkbox": "documentos_pdf_checkbox.png",
    "pantalla Documentos con URL de YouTube detectada": "documentos_youtube.png",
    "pantalla Documentos con un archivo seleccionado": "documentos.png",
    "pantalla Documentos despues de subir": "documentos_subido.png",
    "pantalla Chat en modo Chat": "chat_modo_chat.png",
    "pantalla Chat en modo Query": "chat_modo_query.png",
    "chat PDF con trazabilidad": "chat_apuntes_java.png",
    "chat CSV con trazabilidad": "chat_modo_query.png",
    "chat imagen con trazabilidad": "chat_imagen.png",
    "chat video POO con trazabilidad": "chat_video_poo.png",
    "chat apuntes Java con trazabilidad": "chat_apuntes_java.png",
}


def set_paragraph_font(paragraph, font_name="Atkinson Hyperlegible Next", font_size=11, bold=False, color=None):
    """Aplica fuente a todo el texto de un parrafo."""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color


def add_styled_paragraph(doc, text, font_name="Atkinson Hyperlegible Next", font_size=11, bold=False,
                         color=None, alignment=None, space_after=None, space_before=None):
    """Anade un parrafo con estilo personalizado."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def add_code_block(doc, lines):
    """Anade un bloque de codigo con fuente monoespaciada y fondo gris claro."""
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.left_indent = Cm(1)
        # Fondo gris claro
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
        p._element.get_or_add_pPr().append(shading)


def find_captura_image(captura_text):
    """Busca la imagen correspondiente a una marca [CAPTURA: ...]."""
    for key, filename in CAPTURA_MAP.items():
        if key.lower() in captura_text.lower():
            path = os.path.join(CAPTURAS_DIR, filename)
            if os.path.exists(path):
                return path, filename
    return None, None


def add_image_with_caption(doc, image_path, caption_text):
    """Anade una imagen centrada con pie de foto."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Cm(15))

    # Pie de foto
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = cap.add_run(caption_text)
    run_cap.font.name = "Atkinson Hyperlegible Next"
    run_cap.font.size = Pt(9)
    run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run_cap.font.italic = True
    cap.paragraph_format.space_after = Pt(12)


def is_heading_line(line):
    """Detecta si una linea es un heading y devuelve (nivel, texto)."""
    # Lineas tipo "1. INTRODUCCION", "2. PIPELINE...", etc.
    m = re.match(r'^(\d+)\.\s+(.+)$', line)
    if m:
        section_num = m.group(1)
        title = m.group(2).strip()
        # Si el titulo esta en mayusculas o empieza con mayuscula, es heading 2
        if title == title.upper() or len(title) > 3:
            return 2, f"{section_num}. {title}"

    # Subsecciones tipo "3.1 Extraccion", "5.2 Descripcion..."
    m = re.match(r'^(\d+\.\d+)\s+(.+)$', line)
    if m:
        return 3, f"{m.group(1)} {m.group(2).strip()}"

    # Secciones tipo "CASO 1: CONSULTAS..."
    m = re.match(r'^(CASO \d+:.+)$', line)
    if m:
        return 2, m.group(1)

    # Secciones numeradas del informe tipo "6.1 PANTALLA DE..."
    m = re.match(r'^(\d+\.\d+)\s+(PANTALLA.+|PANTALLA.+)$', line)
    if m:
        return 3, f"{m.group(1)} {m.group(2).strip()}"

    return None, None


def process_text_file(doc, filepath, part_title):
    """Procesa un archivo .txt y lo anade al documento con formato."""
    with open(filepath, "r", encoding="utf-8") as f:
        lines = f.readlines()

    # Saltar las primeras lineas de cabecera del archivo
    start = 0
    for i, line in enumerate(lines):
        if line.strip().startswith("===="):
            start = i + 1
            break
    if start == 0:
        # Si no hay separador, buscar la primera linea en blanco despues de las primeras 3
        for i in range(min(3, len(lines)), len(lines)):
            if not lines[i].strip():
                start = i + 1
                break

    code_block = []
    in_code_block = False

    i = start
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Linea en blanco
        if not stripped:
            if in_code_block and code_block:
                add_code_block(doc, code_block)
                code_block = []
                in_code_block = False
            i += 1
            continue

        # Marca [CAPTURA: ...]
        captura_match = re.match(r'^\[CAPTURA:\s*(.+)\]$', stripped)
        if captura_match:
            if in_code_block and code_block:
                add_code_block(doc, code_block)
                code_block = []
                in_code_block = False
            captura_text = captura_match.group(1)
            img_path, img_name = find_captura_image(captura_text)
            if img_path:
                add_image_with_caption(doc, img_path, captura_text)
            i += 1
            continue

        # Linea de separacion (====, ----)
        if re.match(r'^[=\-]{3,}$', stripped):
            i += 1
            continue

        # Detectar heading
        level, heading_text = is_heading_line(stripped)
        if level and not line.startswith("    ") and not line.startswith("\t"):
            if in_code_block and code_block:
                add_code_block(doc, code_block)
                code_block = []
                in_code_block = False
            h = doc.add_heading(heading_text, level=level)
            set_paragraph_font(h, "Atkinson Hyperlegible Next", 14 if level == 2 else 12, bold=True)
            i += 1
            continue

        # Bloque de codigo (indentado con 4+ espacios)
        if line.startswith("    ") or line.startswith("\t"):
            code_text = line[4:] if line.startswith("    ") else line[1:]
            code_block.append(code_text)
            in_code_block = True
            i += 1
            continue

        # Si estabamos en bloque de codigo, cerrar
        if in_code_block and code_block:
            add_code_block(doc, code_block)
            code_block = []
            in_code_block = False

        # Lista con guion
        if stripped.startswith("- "):
            text = stripped[2:]
            p = doc.add_paragraph(text, style="List Bullet")
            for run in p.runs:
                run.font.name = "Atkinson Hyperlegible Next"
                run.font.size = Pt(11)
            i += 1
            continue

        # Lista con numero y punto (1. 2. 3.)
        m_num = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if m_num and not is_heading_line(stripped)[0]:
            p = doc.add_paragraph(stripped, style="List Number")
            for run in p.runs:
                run.font.name = "Atkinson Hyperlegible Next"
                run.font.size = Pt(11)
            i += 1
            continue

        # Parrafo normal
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.font.name = "Atkinson Hyperlegible Next"
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)
        i += 1

    # Cerrar ultimo bloque de codigo si queda
    if in_code_block and code_block:
        add_code_block(doc, code_block)


def create_cover_page(doc):
    """Crea la portada del documento."""
    # Espaciado superior
    for _ in range(6):
        doc.add_paragraph()

    # Titulo principal
    add_styled_paragraph(doc, "MauiRAG", "Atkinson Hyperlegible Next", 36, bold=True,
                         color=RGBColor(0x51, 0x2B, 0xD4),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_styled_paragraph(doc, "Aplicacion MAUI para RAG local\ncon AnythingLLM y LM Studio",
                         "Atkinson Hyperlegible Next", 16, bold=False,
                         color=RGBColor(0x55, 0x55, 0x55),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    # Linea separadora
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(24)

    # Info del curso
    for text in [
        "Tarea 5 - Sistemas RAG",
        "Programacion de Inteligencia Artificial",
        "CE en Inteligencia Artificial y Big Data",
        "",
        "Febrero 2026",
    ]:
        add_styled_paragraph(doc, text, "Atkinson Hyperlegible Next", 14, bold=False,
                             color=RGBColor(0x33, 0x33, 0x33),
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    doc.add_page_break()


def create_index(doc):
    """Crea el indice del documento."""
    doc.add_heading("INDICE", level=1)

    index_items = [
        ("PARTE I: DOCUMENTO DE PREPROCESADO DE DATOS", [
            "1. Introduccion",
            "2. Pipeline de preprocesado general",
            "3. Ejemplo 1: PDF de historial de empleados",
            "4. Ejemplo 2: CSV de empleados",
            "5. Ejemplo 3: Imagenes (Vision + OCR + EXIF)",
            "6. Ejemplo 4: Transcripcion de video de YouTube",
            "7. Ejemplo 5: Apuntes en PDF segmentados por pagina",
            "8. Conclusiones",
        ]),
        ("PARTE II: INFORME DE USO", [
            "1. Descripcion general",
            "2. Arquitectura MVVM",
            "3. Stack tecnologico",
            "4. Instalacion y puesta en marcha",
            "5. Scripts de preprocesado",
            "6. Guia de usuario",
            "7. Notas de compilacion",
        ]),
        ("PARTE III: EJEMPLOS DE CASOS DE USO", [
            "Caso 1: Consultas sobre el historial de empleados (PDF)",
            "Caso 2: Consultas sobre el CSV de salarios",
            "Caso 3: Consultas sobre imagenes (Vision + OCR + EXIF)",
            "Caso 4: Video educativo (transcripcion de YouTube)",
            "Caso 5: Apuntes de programacion en PDF (segmentados por pagina)",
        ]),
    ]

    for part_title, sections in index_items:
        p = doc.add_paragraph()
        run = p.add_run(part_title)
        run.font.name = "Atkinson Hyperlegible Next"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x51, 0x2B, 0xD4)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

        for section in sections:
            s = doc.add_paragraph()
            run_s = s.add_run(f"    {section}")
            run_s.font.name = "Atkinson Hyperlegible Next"
            run_s.font.size = Pt(11)
            run_s.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            s.paragraph_format.space_after = Pt(2)

    doc.add_page_break()


def main():
    doc = Document()

    # Configurar estilos por defecto
    style = doc.styles["Normal"]
    style.font.name = "Atkinson Hyperlegible Next"
    style.font.size = Pt(11)

    # Margenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 1. Portada
    create_cover_page(doc)

    # 2. Indice
    create_index(doc)

    # 3. PARTE I
    doc.add_heading("PARTE I: DOCUMENTO DE PREPROCESADO DE DATOS", level=1)
    h1 = doc.paragraphs[-1]
    for run in h1.runs:
        run.font.color.rgb = RGBColor(0x51, 0x2B, 0xD4)
    process_text_file(doc, os.path.join(BASE_DIR, "Documento_Preprocesado.txt"), "PARTE I")
    doc.add_page_break()

    # 4. PARTE II
    doc.add_heading("PARTE II: INFORME DE USO", level=1)
    h2 = doc.paragraphs[-1]
    for run in h2.runs:
        run.font.color.rgb = RGBColor(0x51, 0x2B, 0xD4)
    process_text_file(doc, os.path.join(BASE_DIR, "Informe_de_Uso.txt"), "PARTE II")
    doc.add_page_break()

    # 5. PARTE III
    doc.add_heading("PARTE III: EJEMPLOS DE CASOS DE USO", level=1)
    h3 = doc.paragraphs[-1]
    for run in h3.runs:
        run.font.color.rgb = RGBColor(0x51, 0x2B, 0xD4)
    process_text_file(doc, os.path.join(BASE_DIR, "Casos_de_Uso.txt"), "PARTE III")

    # Guardar
    output_path = os.path.join(BASE_DIR, "Informe_MauiRAG.docx")
    doc.save(output_path)
    print(f"Documento generado: {output_path}")
    print(f"Tamano: {os.path.getsize(output_path) / 1024:.0f} KB")


if __name__ == "__main__":
    main()
