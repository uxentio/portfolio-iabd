from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
import re

BASE = r"D:\Tarea 5 RAG"
CAPTURAS = os.path.join(BASE, "Capturas")

# mapeo de placeholders a capturas
CAPTURA_MAP = {
    "pantalla Config": "config.png",
    "pantalla Workspaces": "workspaces.png",
    "pantalla Documentos con un archivo": "documentos_imagen.png",
    "pantalla Documentos despues de subir": "documentos_imagen_ok.png",
    "pantalla Chat en modo Chat": "chat_modo_chat.png",
    "pantalla Chat en modo Query": "chat_modo_query.png",
    "chat PDF": "chat_modo_chat.png",
    "chat CSV": "chat_modo_query.png",
    "chat imagen": "chat_imagen.png",
    "chat video POO": "chat_video_poo.png",
    "chat apuntes Java": "chat_apuntes_java.png",
}

# los 3 documentos fuente, con titulo de seccion y si llevan capturas
DOCUMENTOS = [
    {
        "txt": "Documento_Preprocesado.txt",
        "titulo_seccion": "PARTE I: DOCUMENTO DE PREPROCESADO DE DATOS",
        "capturas": False,
    },
    {
        "txt": "Informe_de_Uso.txt",
        "titulo_seccion": "PARTE II: INFORME DE USO",
        "capturas": True,
    },
    {
        "txt": "Casos_de_Uso.txt",
        "titulo_seccion": "PARTE III: EJEMPLOS DE CASOS DE USO",
        "capturas": True,
    },
]


def agregar_portada(doc):
    """Crea la portada del informe."""
    # espacios antes del titulo para centrarlo verticalmente
    for _ in range(6):
        doc.add_paragraph()

    # titulo principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MauiRAG")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    # subtitulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Aplicacion MAUI para RAG local\ncon AnythingLLM y LM Studio")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    # linea separadora
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    # datos del alumno
    datos = [
        "Tarea 5 - Sistemas RAG",
        "",
        "Programacion de Inteligencia Artificial",
        "CE en Inteligencia Artificial y Big Data",
        "",
        "Febrero 2026",
    ]
    for dato in datos:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if dato == "":
            continue
        run = p.add_run(dato)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(80, 80, 80)

    # salto de pagina despues de la portada
    doc.add_page_break()


def forzar_actualizacion_campos(doc):
    """Anade updateFields al documento para que Word actualice el TOC al abrir."""
    settings = doc.settings.element
    update = parse_xml(f'<w:updateFields {nsdecls("w")} w:val="true"/>')
    settings.append(update)


def agregar_indice(doc):
    """Inserta un campo TOC que Word actualizara al abrir el documento."""
    doc.add_heading("INDICE", level=0)

    # insertar campo TOC con XML directo en el paragraph
    paragraph = doc.add_paragraph()
    p_element = paragraph._p

    # begin
    r1 = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:fldChar w:fldCharType="begin"/>'
        f'</w:r>'
    )
    p_element.append(r1)

    # instruccion TOC
    r2 = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'
        f'</w:r>'
    )
    p_element.append(r2)

    # separate
    r3 = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:fldChar w:fldCharType="separate"/>'
        f'</w:r>'
    )
    p_element.append(r3)

    # texto placeholder (se reemplazara al actualizar)
    r4 = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr><w:color w:val="808080"/></w:rPr>'
        f'  <w:t>Haz clic aqui y pulsa F9 para ver el indice</w:t>'
        f'</w:r>'
    )
    p_element.append(r4)

    # end
    r5 = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:fldChar w:fldCharType="end"/>'
        f'</w:r>'
    )
    p_element.append(r5)

    doc.add_page_break()

    # forzar que Word pregunte si actualizar campos al abrir
    forzar_actualizacion_campos(doc)


def procesar_lineas(doc, lineas, insertar_capturas, offset_nivel=0):
    """Procesa las lineas de un .txt y las anade al documento.
    offset_nivel permite que los heading 1 del txt se conviertan en heading 2 del doc, etc.
    """
    i = 0
    while i < len(lineas):
        linea = lineas[i].rstrip('\n')

        # lineas vacias
        if linea.strip() == '':
            i += 1
            continue

        # placeholders de captura
        if insertar_capturas and linea.strip().startswith('[CAPTURA:'):
            captura_insertada = False
            for clave, archivo in CAPTURA_MAP.items():
                if clave in linea:
                    ruta_img = os.path.join(CAPTURAS, archivo)
                    if os.path.exists(ruta_img):
                        doc.add_picture(ruta_img, width=Inches(5.8))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # pie de captura
                        texto_pie = linea.strip().replace('[CAPTURA: ', '').replace(']', '')
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(texto_pie)
                        run.font.size = Pt(9)
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(128, 128, 128)
                        captura_insertada = True
                        break
            if not captura_insertada:
                doc.add_paragraph(linea)
            i += 1
            continue

        # titulo principal del txt (lineas con === debajo) -> lo salto porque
        # ya pongo el titulo de seccion como heading 1
        if i + 1 < len(lineas) and lineas[i + 1].strip().startswith('===='):
            # no agrego nada, el titulo de parte ya lo cubre
            i += 2
            continue

        # linea de solo === o --- (separador)
        if linea.strip().startswith('====') or linea.strip().startswith('----'):
            i += 1
            continue

        # secciones numeradas principales (1. TITULO) -> heading 2
        match_seccion = re.match(r'^(\d+)\.\s+(.+)', linea.strip())
        if match_seccion and linea.strip() == linea.strip().upper():
            nivel = min(2 + offset_nivel, 3)
            doc.add_heading(linea.strip(), level=nivel)
            i += 1
            continue

        # subsecciones (4.1, 5.2, etc.) -> heading 3
        match_sub = re.match(r'^(\d+\.\d+)\s+(.+)', linea.strip())
        if match_sub:
            nivel = min(3 + offset_nivel, 4)
            doc.add_heading(linea.strip(), level=nivel)
            i += 1
            continue

        # secciones CASO X: -> heading 2
        if linea.strip().startswith('CASO ') and ':' in linea:
            nivel = min(2 + offset_nivel, 3)
            doc.add_heading(linea.strip(), level=nivel)
            i += 1
            continue

        # secciones con --- alrededor
        if linea.strip().startswith('---') and linea.strip().endswith('---') and len(linea.strip()) > 6:
            titulo = linea.strip().strip('-').strip()
            if titulo:
                nivel = min(3 + offset_nivel, 4)
                doc.add_heading(titulo, level=nivel)
            i += 1
            continue

        # bloques indentados (codigo o ejemplos)
        if linea.startswith('    ') and not linea.strip().startswith('-') and not linea.strip().startswith('*'):
            bloque = []
            while i < len(lineas) and (lineas[i].startswith('    ') or lineas[i].strip() == ''):
                if lineas[i].strip() == '' and (i + 1 >= len(lineas) or not lineas[i + 1].startswith('    ')):
                    break
                bloque.append(lineas[i].rstrip('\n'))
                i += 1
            texto_bloque = '\n'.join(bloque)
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(texto_bloque.strip())
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            continue

        # listas con - o *
        if linea.strip().startswith('- ') or linea.strip().startswith('* '):
            texto_item = linea.strip()[2:]
            doc.add_paragraph(texto_item, style='List Bullet')
            i += 1
            continue

        # listas numeradas (1. , 2. , etc.) que NO son titulos en mayusculas
        match_lista_num = re.match(r'^(\d+)\.\s+(.+)', linea.strip())
        if match_lista_num and linea.strip() != linea.strip().upper():
            doc.add_paragraph(linea.strip(), style='List Number')
            i += 1
            continue

        # parrafo normal
        doc.add_paragraph(linea.strip())
        i += 1


def crear_informe_unificado():
    """Genera un unico .docx con portada, indice y los 3 documentos."""
    doc = Document()

    # estilos base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # margenes del documento
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 1. portada
    agregar_portada(doc)

    # 2. indice
    agregar_indice(doc)

    # 3. cada documento como una parte del informe
    for info_doc in DOCUMENTOS:
        txt_path = os.path.join(BASE, info_doc["txt"])

        if not os.path.exists(txt_path):
            print(f"AVISO: No se encuentra {txt_path}, saltando...")
            continue

        # titulo de la parte (Heading 1)
        doc.add_heading(info_doc["titulo_seccion"], level=1)

        with open(txt_path, 'r', encoding='utf-8') as f:
            lineas = f.readlines()

        procesar_lineas(doc, lineas, info_doc["capturas"], offset_nivel=0)

        # salto de pagina entre partes (excepto la ultima)
        if info_doc != DOCUMENTOS[-1]:
            doc.add_page_break()

    # guardar
    ruta_salida = os.path.join(BASE, "Informe_MauiRAG.docx")
    doc.save(ruta_salida)
    print(f"Generado: {ruta_salida}")
    return ruta_salida


def crear_docx_individual(txt_path, docx_path, insertar_capturas=False):
    """Genera un .docx individual a partir de un .txt (por si se necesitan sueltos)."""
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(txt_path, 'r', encoding='utf-8') as f:
        lineas = f.readlines()

    i = 0
    while i < len(lineas):
        linea = lineas[i].rstrip('\n')

        if linea.strip() == '':
            i += 1
            continue

        if insertar_capturas and linea.strip().startswith('[CAPTURA:'):
            captura_insertada = False
            for clave, archivo in CAPTURA_MAP.items():
                if clave in linea:
                    ruta_img = os.path.join(CAPTURAS, archivo)
                    if os.path.exists(ruta_img):
                        doc.add_picture(ruta_img, width=Inches(5.8))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        texto_pie = linea.strip().replace('[CAPTURA: ', '').replace(']', '')
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(texto_pie)
                        run.font.size = Pt(9)
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(128, 128, 128)
                        captura_insertada = True
                        break
            if not captura_insertada:
                doc.add_paragraph(linea)
            i += 1
            continue

        if i + 1 < len(lineas) and lineas[i + 1].strip().startswith('===='):
            doc.add_heading(linea.strip(), level=0)
            i += 2
            continue

        if linea.strip().startswith('====') or linea.strip().startswith('----'):
            i += 1
            continue

        match_seccion = re.match(r'^(\d+)\.\s+(.+)', linea.strip())
        if match_seccion and linea.strip() == linea.strip().upper():
            doc.add_heading(linea.strip(), level=1)
            i += 1
            continue

        match_sub = re.match(r'^(\d+\.\d+)\s+(.+)', linea.strip())
        if match_sub:
            doc.add_heading(linea.strip(), level=2)
            i += 1
            continue

        if linea.strip().startswith('CASO ') and ':' in linea:
            doc.add_heading(linea.strip(), level=1)
            i += 1
            continue

        if linea.strip().startswith('---') and linea.strip().endswith('---') and len(linea.strip()) > 6:
            titulo = linea.strip().strip('-').strip()
            if titulo:
                doc.add_heading(titulo, level=2)
            i += 1
            continue

        if linea.startswith('    ') and not linea.strip().startswith('-') and not linea.strip().startswith('*'):
            bloque = []
            while i < len(lineas) and (lineas[i].startswith('    ') or lineas[i].strip() == ''):
                if lineas[i].strip() == '' and (i + 1 >= len(lineas) or not lineas[i + 1].startswith('    ')):
                    break
                bloque.append(lineas[i].rstrip('\n'))
                i += 1
            texto_bloque = '\n'.join(bloque)
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(texto_bloque.strip())
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            continue

        if linea.strip().startswith('- ') or linea.strip().startswith('* '):
            texto_item = linea.strip()[2:]
            doc.add_paragraph(texto_item, style='List Bullet')
            i += 1
            continue

        doc.add_paragraph(linea.strip())
        i += 1

    doc.save(docx_path)
    print(f"Generado: {docx_path}")


# === GENERAR ===

# 1. informe unificado con portada e indice
print("=== Generando informe unificado ===")
crear_informe_unificado()

# 2. documentos individuales (por si los necesitas sueltos)
print("\n=== Generando documentos individuales ===")
crear_docx_individual(
    os.path.join(BASE, "Documento_Preprocesado.txt"),
    os.path.join(BASE, "Documento_Preprocesado.docx"),
    insertar_capturas=False
)

crear_docx_individual(
    os.path.join(BASE, "Informe_de_Uso.txt"),
    os.path.join(BASE, "Informe_de_Uso.docx"),
    insertar_capturas=True
)

crear_docx_individual(
    os.path.join(BASE, "Casos_de_Uso.txt"),
    os.path.join(BASE, "Casos_de_Uso.docx"),
    insertar_capturas=True
)

print("\nTodos los documentos generados correctamente!")
print("  - Informe_MauiRAG.docx (unificado con portada e indice)")
print("  - Documento_Preprocesado.docx")
print("  - Informe_de_Uso.docx")
print("  - Casos_de_Uso.docx")
