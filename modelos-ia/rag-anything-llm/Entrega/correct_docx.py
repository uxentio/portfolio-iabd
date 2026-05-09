import docx
import re
import xml.etree.ElementTree as ET

doc_path = "Informe_MauiRAG.docx"
doc = docx.Document(doc_path)

replacements = {
    r"\bcomo he preparado\b": "cómo he preparado",
    r"\bque pasa\b": "qué pasa",
    r"\bpequenos\b": "pequeños",
    r"\btendran\b": "tendrán",
    r"\bmas similares\b": "más similares",
    r"\bQuien tiene mas\b": "Quién tiene más",
    r"\bhabria\b": "habría",
    r"\blinea a linea\b": "línea a línea",
    r"\blinea\b": "línea",
    r"\bConsideraciónes\b": "Consideraciones",
    r"\bentre si\b": "entre sí",
    r"\bTambien\b": "También",
    r"\btambien\b": "también",
    r"\basegure\b": "aseguré",
    r"\btenia\b": "tenía",
    r"\bque representa\b": "qué representa",
    r"\bfunciónalidades\b": "funcionalidades",
    r"\btecnicas\b": "técnicas",
    r"\bseleccióna\b": "selecciona",
    r"\benvia\b": "envía",
    r"\bmas aporta\b": "más aporta",
    r"\bfuncióne\b": "funcione",
    r"\banalisis\b": "análisis",
    r"\bfunciónando\b": "funcionando",
    r"\btamano\b": "tamaño",
    r"\bUltima\b": "Última",
    r"\bdemas\b": "demás",
    r"\bvehiculo\b": "vehículo",
    r"\bsubio\b": "subió",
    r"\bCuando se hizo\b": "Cuándo se hizo",
    r"\bQue texto\b": "Qué texto",
    r"\bQue marca\b": "Qué marca",
    r"\basi que\b": "así que",
    r"\bmayoria\b": "mayoría",
    r"\bsubtitulos\b": "subtítulos",
    r"\bautomaticos\b": "automáticos",
    r"\ba traves de\b": "a través de",
    r"\bunico\b": "único",
    r"\bespanol\b": "español",
    r"\bpestana\b": "pestaña",
    r"\bboton\b": "botón",
    r"\bpodria\b": "podría",
    r"\bQue es\b": "Qué es",
    r"\bseleccióne\b": "seleccioné",
    r"\banade\b": "añade",
    r"\bfunciónaba\b": "funcionaba",
    r"\bdescubri\b": "descubrí",
    r"\bdevolvia\b": "devolvía",
    r"\bcomodo\b": "cómodo",
    r"\bseríalización\b": "serialización",
    r"\bdeseríalización\b": "deserialización",
    r"\bfuncióna\b": "funciona",
    r"\bprobe\b": "probé",
    r"\badiciónal\b": "adicional",
    r"\bopciónal\b": "opcional",
    r"\bSelecciónar\b": "Seleccionar",
    r"\banadan\b": "añadan",
    r"\bGuia\b": "Guía",
    r"\bguia\b": "guía",
    r"\bgráfica\b": "gráfica",
    r"\bfunción\b": "función",
    }

def replace_in_runs(runs):
    # Quick fix: replace within each run.
    # Note: If a word is split across runs, it might not be fixed. We join and split if needed, 
    # but the simplest is just run.text.replace.
    for run in runs:
        for old, new in replacements.items():
            run.text = re.sub(old, new, run.text)

for p in doc.paragraphs:
    replace_in_runs(p.runs)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_runs(p.runs)

doc.save("Informe_MauiRAG.docx")
print("Done modifying Informe_MauiRAG.docx")
