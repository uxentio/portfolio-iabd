"""
Script para generar el documento Word de la Tarea 7
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# ---------------------------------------------------
# Intentar crear diagrama de arquitectura con matplotlib
# ---------------------------------------------------
diagrama_path = None
try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.patches as patches

    fig, ax = plt.subplots(1, 1, figsize=(10, 5.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5.5)
    ax.axis('off')

    # Cajas del diagrama
    cajas = [
        (0.3, 2.2, 1.6, 1.1, u'C\u00e1mara /\nGaler\u00eda', '#E8F5E9', '#388E3C'),
        (2.6, 2.2, 2.0, 1.1, 'App .NET MAUI\n(FacturasOCR)', '#E3F2FD', '#1565C0'),
        (5.6, 2.2, 2.3, 1.1, 'Azure Document\nIntelligence', '#FFF3E0', '#E65100'),
        (2.6, 0.4, 1.2, 0.9, 'SQLite\n(local)', '#F3E5F5', '#7B1FA2'),
        (4.2, 0.4, 1.3, 0.9, 'Backup\n(opcional)', '#F5F5F5', '#757575'),
        (8.4, 2.2, 1.3, 1.1, 'Resultados\nOCR', '#E8F5E9', '#2E7D32'),
    ]

    for (x, y, w, h, text, bg, border) in cajas:
        rect = patches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.08",
                                       facecolor=bg, edgecolor=border, linewidth=2)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, text, ha='center', va='center',
                fontsize=9, fontweight='bold', color='#333333')

    # Flechas
    props = dict(arrowstyle='->', color='#444', lw=2, connectionstyle='arc3,rad=0')
    ax.annotate('', xy=(2.6, 2.75), xytext=(1.9, 2.75), arrowprops=props)
    ax.annotate('', xy=(5.6, 2.9), xytext=(4.6, 2.9), arrowprops=props)
    ax.annotate('', xy=(4.6, 2.5), xytext=(5.6, 2.5), arrowprops=dict(arrowstyle='->', color='#444', lw=2))
    ax.annotate('', xy=(8.4, 2.75), xytext=(7.9, 2.75), arrowprops=props)
    ax.annotate('', xy=(3.2, 2.2), xytext=(3.2, 1.3), arrowprops=props)
    ax.annotate('', xy=(4.8, 2.2), xytext=(4.8, 1.3), arrowprops=props)

    # Textos en flechas
    ax.text(2.25, 3.05, 'imagen/PDF', fontsize=7, ha='center', color='#666', style='italic')
    ax.text(5.1, 3.15, 'stream HTTPS', fontsize=7, ha='center', color='#666', style='italic')
    ax.text(5.1, 2.2, 'AnalyzeResult', fontsize=7, ha='center', color='#666', style='italic')
    ax.text(8.15, 3.0, 'mostrar', fontsize=7, ha='center', color='#666', style='italic')
    ax.text(2.8, 1.65, 'guardar', fontsize=7, ha='center', color='#666', style='italic')
    ax.text(5.3, 1.65, 'copiar', fontsize=7, ha='center', color='#666', style='italic')

    ax.text(5, 4.8, u'Arquitectura - FacturasOCR', ha='center', fontsize=14, fontweight='bold', color='#222')

    diagrama_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '_temp_diagrama.png')
    plt.savefig(diagrama_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    print("Diagrama generado OK")
except Exception as ex:
    print(f"No se pudo generar el diagrama: {ex}")
    diagrama_path = None

# ---------------------------------------------------
# Crear documento Word
# ---------------------------------------------------
doc = Document()

# Estilo base
style = doc.styles['Normal']
style.font.name = 'Atkinson Hyperlegible Next'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Headings con la misma fuente
for i in range(1, 4):
    h_style = doc.styles[f'Heading {i}']
    h_style.font.name = 'Atkinson Hyperlegible Next'

FUENTE = 'Atkinson Hyperlegible Next'

def aplicar_fuente_tabla(table, fuente=FUENTE):
    """Aplica la fuente a todas las celdas de una tabla (no heredan de Normal)."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.name is None or run.font.name == 'Calibri':
                        run.font.name = fuente

# Margenes
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============== PORTADA ==============
for _ in range(5):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(u'Tarea 7 \u2013 Servicios Cognitivos de Azure')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(u'Proyecto 2: OCR de Facturas con\nAzure Document Intelligence')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(80, 80, 80)

for _ in range(3):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(u'autor Mart\u00ednez')
run.font.size = Pt(14)
run.font.bold = True

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Marzo 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ============== INDICE ==============
h = doc.add_heading(u'\u00cdndice', level=1)

indice_items = [
    ('1. Introducci\u00f3n', '3'),
    ('2. Arquitectura general', '4'),
    ('3. Servicios de Azure utilizados', '5'),
    ('    3.1. Azure Document Intelligence', '5'),
    ('    3.2. Configuraci\u00f3n del recurso', '5'),
    ('    3.3. Paquete NuGet y c\u00f3digo', '6'),
    ('    3.4. Estructura del AnalyzeResult', '7'),
    ('    3.5. Diferencias con el ejemplo de clase', '8'),
    ('4. Estructura del proyecto', '9'),
    ('    4.1. Modelo de datos (ResultadoFactura)', '9'),
    ('    4.2. Servicio de base de datos', '10'),
    ('5. Requisitos e instalaci\u00f3n', '11'),
    ('    5.1. Requisitos previos', '11'),
    ('    5.2. C\u00f3mo ejecutar el proyecto', '11'),
    ('    5.3. Configuraci\u00f3n de Azure', '12'),
    ('6. Funcionalidades principales', '13'),
    ('    6.1. Procesamiento de facturas', '13'),
    ('    6.2. Procesamiento por lotes', '14'),
    ('    6.3. Historial y detalle', '14'),
    ('    6.4. Multiidioma', '14'),
    ('    6.5. Tema claro / oscuro', '15'),
    ('7. Capturas de pantalla', '16'),
    ('8. Conclusiones', '18'),
]

for titulo, pagina in indice_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = p.add_run(titulo)
    run.font.size = Pt(11)
    run_page = p.add_run(f'\t{pagina}')
    run_page.font.size = Pt(11)
    run_page.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()


# ============== 1. INTRODUCCION ==============
doc.add_heading(u'1. Introducci\u00f3n', level=1)

doc.add_paragraph(
    u'He elegido el Proyecto 2 de la pr\u00e1ctica (OCR de Facturas con Azure Document Intelligence). '
    u'La idea es desarrollar una app .NET MAUI que permita hacer una foto a una factura con la '
    u'c\u00e1mara del m\u00f3vil o elegir un archivo (im\u00e1genes, PDFs, XMLs de factura electr\u00f3nica) '
    u'y que la app mande ese archivo al servicio de Azure Document Intelligence para extraer '
    u'los datos importantes: nombre del proveedor, n\u00famero de factura, fecha e importe total.'
)

doc.add_paragraph(
    u'Para desarrollar la app me he basado en los materiales de clase, sobre todo en el documento '
    u'"Creaci\u00f3n del recurso Form Recognizer" donde viene explicado paso a paso c\u00f3mo crear el '
    u'recurso en Azure y el c\u00f3digo de ejemplo en C# con el paquete Azure.AI.FormRecognizer. '
    u'Tambi\u00e9n he consultado la presentaci\u00f3n "Azure y Servicios de IA en la Nube" para entender '
    u'las distintas categor\u00edas de servicios cognitivos y el glosario de t\u00e9rminos de Azure '
    u'para tener claros los conceptos b\u00e1sicos (recursos, grupos de recursos, suscripciones, etc.).'
)

doc.add_paragraph(
    u'Los datos extra\u00eddos se guardan autom\u00e1ticamente en una base de datos SQLite dentro de la app, '
    u'y se pueden consultar despu\u00e9s en la pesta\u00f1a de historial. Tambi\u00e9n he a\u00f1adido un sistema '
    u'de copia de seguridad opcional por si el usuario quiere tener los archivos guardados en otra carpeta.'
)

doc.add_paragraph(
    u'La app est\u00e1 disponible en 9 idiomas (castellano, ingl\u00e9s, y las lenguas cooficiales '
    u'de Espa\u00f1a: catal\u00e1n, euskera, gallego, aragon\u00e9s, aran\u00e9s, asturiano y extreme\u00f1o) '
    u'y se puede cambiar el idioma en tiempo real desde los ajustes sin tener que reiniciar la app.'
)


# ============== 2. ARQUITECTURA ==============
doc.add_page_break()
doc.add_heading('2. Arquitectura general', level=1)

doc.add_paragraph(
    u'La arquitectura de la aplicaci\u00f3n es relativamente sencilla. El usuario interact\u00faa '
    u'con la interfaz MAUI (que tiene tres pesta\u00f1as: Analizar, Archivo y Ajustes), '
    u'la app manda las im\u00e1genes o PDFs al servicio de Azure por HTTPS, y cuando Azure responde '
    u'con los datos extra\u00eddos, la app los formatea, los muestra en pantalla y los guarda en '
    u'una base de datos SQLite local. Opcionalmente tambi\u00e9n puede copiar los archivos a una '
    u'carpeta de backup que elija el usuario.'
)

# Insertar diagrama
if diagrama_path and os.path.exists(diagrama_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(diagrama_path, width=Inches(5.8))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run('Fig. 1: Esquema de arquitectura de la app')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
else:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[ DIAGRAMA DE ARQUITECTURA - insertar manualmente ]')
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.font.italic = True

doc.add_paragraph('')

doc.add_paragraph(
    u'El flujo de procesamiento es el siguiente:'
)

pasos = [
    u'El usuario elige una imagen (c\u00e1mara o galer\u00eda) o selecciona un archivo PDF/XML.',
    u'Si es una imagen o PDF, la app manda el archivo al servicio de Azure Document Intelligence por HTTPS usando el modelo prebuilt-invoice.',
    u'Azure analiza el documento y devuelve un objeto AnalyzeResult con pares clave-valor, campos espec\u00edficos de factura (VendorName, InvoiceTotal...), p\u00e1ginas con l\u00edneas de texto y tablas.',
    u'La app formatea el resultado, extrae los campos principales y los muestra en una ficha visual.',
    u'Se guarda todo en la base de datos SQLite local y opcionalmente se copia a la carpeta de backup.',
    u'Si el archivo es XML (factura electr\u00f3nica FacturaE), se parsea localmente sin necesidad de Azure.',
]

for i, paso in enumerate(pasos):
    doc.add_paragraph(f'{i+1}. {paso}')


# ============== 3. SERVICIOS AZURE ==============
doc.add_page_break()
doc.add_heading('3. Servicios de Azure utilizados', level=1)

doc.add_heading('3.1. Azure Document Intelligence (Form Recognizer)', level=2)

doc.add_paragraph(
    u'Azure Document Intelligence (antes se llamaba Form Recognizer, como aparece en los apuntes '
    u'de la Unidad 8) es un servicio cognitivo de Azure que usa modelos de IA para extraer texto, '
    u'pares clave-valor, tablas y campos estructurados de documentos. Como se explica en el '
    u'glosario de Azure de clase, es un servicio dentro de la categor\u00eda "Vision" de los '
    u'Azure Cognitive Services. Es como un OCR pero m\u00e1s avanzado porque no solo lee el texto '
    u'sino que entiende la estructura del documento y sabe identificar qu\u00e9 informaci\u00f3n es cada cosa.'
)

doc.add_paragraph(
    u'En mi app uso el modelo prebuilt-invoice, que es un modelo pre-entrenado espec\u00edfico '
    u'para facturas. Este modelo sabe buscar autom\u00e1ticamente campos como:'
)

campos_azure = [
    ('VendorName', 'Nombre del proveedor/emisor de la factura'),
    ('InvoiceId', u'N\u00famero de factura'),
    ('InvoiceDate', 'Fecha de la factura'),
    ('InvoiceTotal', 'Importe total'),
]

table = doc.add_table(rows=len(campos_azure)+1, cols=2, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Campo Azure'
hdr[1].text = u'Descripci\u00f3n'
for cell in table.rows[0].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

for i, (campo, desc) in enumerate(campos_azure):
    cells = table.rows[i+1].cells
    cells[0].text = campo
    cells[0].paragraphs[0].runs[0].font.name = 'Consolas'
    cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    cells[1].text = desc

aplicar_fuente_tabla(table)
doc.add_paragraph('')

doc.add_paragraph(
    u'Adem\u00e1s de estos campos, el resultado de Azure tambi\u00e9n incluye los pares clave-valor '
    u'gen\u00e9ricos (como se muestra en el ejemplo ShowKeyValuePairsSummary del c\u00f3digo de clase), '
    u'las l\u00edneas de texto de cada p\u00e1gina (ShowLayoutAndMarksSummary) y las tablas con sus '
    u'celdas (ShowTablesSummary). Esto es \u00fatil porque no todas las facturas tienen el mismo '
    u'formato y a veces Azure no detecta un campo pero s\u00ed lo tiene en los pares clave-valor.'
)


doc.add_heading(u'3.2. Configuraci\u00f3n del recurso', level=2)

doc.add_paragraph(
    u'Para crear el recurso en Azure segu\u00ed los pasos del documento "Creaci\u00f3n del recurso '
    u'Form Recognizer" que nos proporcion\u00f3 el profesor: Datos B\u00e1sicos, Red (dej\u00e9 el acceso '
    u'p\u00fablico porque es un entorno de pruebas), Identity (desactivado), Tags y Revisar y Crear. '
    u'Us\u00e9 el tier gratuito F0 que permite hasta 500 p\u00e1ginas al mes. '
    u'Una vez creado, en la secci\u00f3n "Claves y punto de conexi\u00f3n" copi\u00e9 el Endpoint y la API Key.'
)

table2 = doc.add_table(rows=4, cols=2, style='Table Grid')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table2.rows[0].cells
hdr[0].text = u'Par\u00e1metro'
hdr[1].text = 'Valor'
for cell in table2.rows[0].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

data2 = [
    ('Endpoint', 'https://uxentio.cognitiveservices.azure.com/'),
    ('Modelo', 'prebuilt-invoice'),
    ('Tier', 'F0 (gratuito)'),
]
for i, (k, v) in enumerate(data2):
    cells = table2.rows[i+1].cells
    cells[0].text = k
    cells[1].text = v

aplicar_fuente_tabla(table2)
doc.add_paragraph('')

doc.add_heading(u'3.3. Paquete NuGet y c\u00f3digo de conexi\u00f3n', level=2)

doc.add_paragraph(
    u'La conexi\u00f3n con Azure se hace con el paquete NuGet Azure.AI.FormRecognizer versi\u00f3n 4.1.0, '
    u'el mismo que se usa en el ejemplo de clase (FormRecognizer/Program.cs). '
    u'El patr\u00f3n es el mismo: crear un DocumentAnalysisClient con el endpoint y la clave, '
    u'llamar a AnalyzeDocumentAsync y procesar el AnalyzeResult. La diferencia es que en el ejemplo '
    u'de clase se usa "prebuilt-document" (gen\u00e9rico) y yo uso "prebuilt-invoice" que est\u00e1 '
    u'optimizado para facturas. El c\u00f3digo b\u00e1sico:'
)

# Code block
code_text = (
    'var client = new DocumentAnalysisClient(\n'
    '    new Uri(endpoint),\n'
    '    new AzureKeyCredential(apiKey));\n'
    '\n'
    'var operation = await client.AnalyzeDocumentAsync(\n'
    '    WaitUntil.Completed,\n'
    '    "prebuilt-invoice",\n'
    '    imageStream);\n'
    '\n'
    'AnalyzeResult result = operation.Value;'
)

# Poner el codigo en un parrafo con fondo gris
p_code = doc.add_paragraph()
p_code.paragraph_format.left_indent = Cm(1)
p_code.paragraph_format.right_indent = Cm(1)
run = p_code.add_run(code_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(30, 30, 30)

# shading for code block
from docx.oxml import OxmlElement
pPr = p_code._element.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'F2F2F2')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
pPr.append(shd)

doc.add_paragraph(
    u'El endpoint y la API key se pueden cambiar desde la pantalla de Ajustes de la app, '
    u'por si el usuario quiere usar sus propias credenciales de Azure.'
)


doc.add_heading(u'3.4. Estructura del AnalyzeResult', level=2)

doc.add_paragraph(
    u'Cuando Azure termina de analizar un documento, devuelve un objeto AnalyzeResult que contiene '
    u'toda la informaci\u00f3n extra\u00edda. Es exactamente lo mismo que se ve en el ejemplo de clase '
    u'(FormRecognizer/Program.cs), donde cada parte se muestra con una funci\u00f3n Show diferente. '
    u'En mi app proceso las mismas partes pero adaptadas a MAUI:'
)

analyze_parts = [
    ('KeyValuePairs', u'Pares clave-valor gen\u00e9ricos', u'Detecta etiquetas y sus valores por proximidad en el documento. Ej: "Fecha" \u2192 "15/01/2024". Equivale a ShowKeyValuePairsSummary del ejemplo.'),
    ('Documents[].Fields', 'Campos del modelo', u'Campos sem\u00e1nticos espec\u00edficos del modelo prebuilt-invoice: VendorName, InvoiceId, InvoiceDate, InvoiceTotal, etc. En el ejemplo de clase (que usa prebuilt-document) esta parte se muestra con ShowEntitiesSummary.'),
    ('Pages[].Lines', u'L\u00edneas de texto por p\u00e1gina', u'Todas las l\u00edneas de texto detectadas, con sus coordenadas (BoundingPolygon). En el ejemplo se muestra con ShowLayoutAndMarksSummary.'),
    ('Pages[].Words', 'Palabras individuales', u'Cada palabra con su texto y nivel de confianza. \u00datil para entender qu\u00e9 ha le\u00eddo Azure exactamente.'),
    ('Tables[].Cells', 'Tablas con celdas', u'Tablas detectadas en el documento, con fila, columna y contenido de cada celda. En el ejemplo se usa ShowTablesSummary. Yo las uso como fallback para buscar totales.'),
]

table_az = doc.add_table(rows=len(analyze_parts)+1, cols=3, style='Table Grid')
table_az.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, header in enumerate(['Propiedad', 'Tipo de datos', u'Descripci\u00f3n y equivalencia en ejemplo de clase']):
    cell = table_az.rows[0].cells[j]
    cell.text = header
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(9)

for i, (prop, tipo, desc) in enumerate(analyze_parts):
    cells = table_az.rows[i+1].cells
    cells[0].text = prop
    for run in cells[0].paragraphs[0].runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
    cells[1].text = tipo
    for run in cells[1].paragraphs[0].runs:
        run.font.size = Pt(9)
    cells[2].text = desc
    for run in cells[2].paragraphs[0].runs:
        run.font.size = Pt(9)

aplicar_fuente_tabla(table_az)
doc.add_paragraph('')

doc.add_paragraph(
    u'En el ejemplo de clase tambi\u00e9n se procesan las secciones manuscritas '
    u'(ShowHandwrittenSectionsSummary) y las marcas de selecci\u00f3n (SelectionMarks), pero en mi app '
    u'no las uso porque las facturas normales no suelen tener texto manuscrito ni checkboxes.'
)


doc.add_heading(u'3.5. Diferencias con el ejemplo de clase', level=2)

doc.add_paragraph(
    u'El c\u00f3digo de ejemplo del profesor (FormRecognizer/Program.cs) y mi app comparten la misma '
    u'base, pero hay varias diferencias importantes que tuve que ir resolviendo:'
)

diferencias = [
    (u'Modelo utilizado', u'El ejemplo usa "prebuilt-document" (gen\u00e9rico, para cualquier documento). '
     u'Yo uso "prebuilt-invoice" que est\u00e1 optimizado para facturas y devuelve campos como '
     u'VendorName, InvoiceTotal, etc. directamente en Documents[].Fields.'),
    (u'Origen del documento', u'El ejemplo analiza un PDF desde una URL (AnalyzeDocumentFromUriAsync). '
     u'Mi app usa un stream local (AnalyzeDocumentAsync) porque la imagen viene de la c\u00e1mara o '
     u'del sistema de archivos del m\u00f3vil.'),
    (u'Formato de salida', u'El ejemplo imprime todo por consola. Mi app formatea el resultado en un '
     u'editor de texto y extrae los campos clave para mostrarlos en una ficha visual con colores.'),
    (u'Plataforma', u'El ejemplo es una app de consola .NET. Mi app es .NET MAUI con interfaz gr\u00e1fica, '
     u'navegaci\u00f3n por pesta\u00f1as, base de datos SQLite y soporte multiidioma.'),
    (u'Acceso a campos', u'En el ejemplo se usa field.Value?.ToString() que funciona bien con prebuilt-document. '
     u'Con prebuilt-invoice tuve que usar field.Content porque los campos tienen tipos ricos '
     u'(Currency, Date) y ToString() no devuelve el valor legible.'),
]

for titulo_dif, desc_dif in diferencias:
    p = doc.add_paragraph()
    run_bold = p.add_run(f'\u2022 {titulo_dif}: ')
    run_bold.font.bold = True
    run_bold.font.size = Pt(10)
    run_normal = p.add_run(desc_dif)
    run_normal.font.size = Pt(10)

doc.add_paragraph('')


# ============== 4. ESTRUCTURA DEL PROYECTO ==============
doc.add_page_break()
doc.add_heading('4. Estructura del proyecto', level=1)

doc.add_paragraph(
    u'El proyecto sigue la estructura t\u00edpica de una app .NET MAUI, con los archivos organizados '
    u'en carpetas seg\u00fan su funci\u00f3n. He puesto los nombres en espa\u00f1ol para que sea m\u00e1s f\u00e1cil de entender:'
)

files_data = [
    ('MauiProgram.cs', 'Punto de entrada, configura fuentes y servicios'),
    ('App.xaml / App.xaml.cs', u'Aplicaci\u00f3n principal, aplica el tema guardado al arrancar'),
    ('AppShell.xaml / .cs', u'Navegaci\u00f3n por pesta\u00f1as (Analizar, Archivo, Ajustes)'),
    ('PaginaPrincipal.xaml / .cs', u'Pantalla principal: botones, preview de imagen, ficha de datos, procesamiento OCR'),
    ('Paginas/PaginaHistorial.xaml / .cs', 'Historial de facturas guardadas con lista'),
    ('Paginas/PaginaDetalle.xaml / .cs', 'Detalle de una factura (imagen, campos, texto completo)'),
    ('Paginas/PaginaAjustes.xaml / .cs', u'Ajustes: Azure, apariencia (tema), idioma, almacenamiento, backup'),
    ('Modelos/ResultadoFactura.cs', 'Modelo de datos / tabla SQLite con campos de factura'),
    ('Servicios/ServicioAnalisis.cs', u'Conexi\u00f3n con Azure Document Intelligence'),
    ('Servicios/ServicioBaseDatos.cs', 'Servicio de base de datos SQLite (CRUD)'),
    ('Servicios/ServicioIdiomas.cs', u'Sistema de traducciones para 9 idiomas'),
]

table3 = doc.add_table(rows=len(files_data)+1, cols=2, style='Table Grid')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table3.rows[0].cells
hdr[0].text = 'Archivo'
hdr[1].text = u'Descripci\u00f3n'
for cell in table3.rows[0].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

for i, (f, d) in enumerate(files_data):
    cells = table3.rows[i+1].cells
    cells[0].text = f
    for run in cells[0].paragraphs[0].runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
    cells[1].text = d
    for run in cells[1].paragraphs[0].runs:
        run.font.size = Pt(10)

aplicar_fuente_tabla(table3)
doc.add_paragraph('')

doc.add_paragraph(
    u'El archivo m\u00e1s grande e importante es PaginaPrincipal.xaml.cs, que contiene toda la l\u00f3gica '
    u'de procesamiento de facturas: llamada a Azure, formateo del resultado, extracci\u00f3n de campos, '
    u'parseo de XML, guardado autom\u00e1tico y detecci\u00f3n de duplicados.'
)

doc.add_heading('4.1. Modelo de datos (ResultadoFactura)', level=2)

doc.add_paragraph(
    u'Cada factura procesada se guarda en una base de datos SQLite local usando la clase '
    u'ResultadoFactura (en Modelos/ResultadoFactura.cs). Esta clase tiene un atributo '
    u'[Table("InvoiceResult")] para que la tabla de la BD se llame siempre igual aunque la '
    u'clase se haya renombrado a espa\u00f1ol. Los campos son:'
)

modelo_campos = [
    ('Id', 'int', u'Clave primaria autoincremental. SQLite se encarga de asignarla.'),
    ('AnalyzedAt', 'DateTime', u'Fecha y hora en que se proces\u00f3 la factura.'),
    ('ImagePath', 'string', u'Ruta donde se guard\u00f3 la copia local del archivo original (imagen, PDF o XML).'),
    ('RawText', 'string', u'Texto completo del OCR: todos los pares clave-valor, campos, l\u00edneas y tablas que devolvi\u00f3 Azure formateados como texto plano.'),
    ('FileName', 'string', 'Nombre del archivo original (ej: "factura_mayo.pdf").'),
    ('VendorName', 'string', u'Nombre del proveedor extra\u00eddo por Azure o por el fallback manual.'),
    ('InvoiceId', 'string', u'N\u00famero de factura extra\u00eddo.'),
    ('InvoiceDate', 'string', u'Fecha de la factura extra\u00edda.'),
    ('Total', 'string', u'Importe total extra\u00eddo.'),
]

table_modelo = doc.add_table(rows=len(modelo_campos)+1, cols=3, style='Table Grid')
table_modelo.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, header in enumerate(['Campo', 'Tipo', u'Descripci\u00f3n']):
    cell = table_modelo.rows[0].cells[j]
    cell.text = header
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(9)

for i, (campo, tipo, desc) in enumerate(modelo_campos):
    cells = table_modelo.rows[i+1].cells
    cells[0].text = campo
    for run in cells[0].paragraphs[0].runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    cells[1].text = tipo
    for run in cells[1].paragraphs[0].runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    cells[2].text = desc
    for run in cells[2].paragraphs[0].runs:
        run.font.size = Pt(9)

aplicar_fuente_tabla(table_modelo)
doc.add_paragraph('')

doc.add_paragraph(
    u'Los campos VendorName, InvoiceId, InvoiceDate y Total se guardan como string y no como '
    u'sus tipos nativos (DateTime, decimal...) porque Azure no siempre los devuelve en un formato '
    u'est\u00e1ndar. Por ejemplo, una fecha puede venir como "15/01/2024", "January 15, 2024" o '
    u'"2024-01-15" seg\u00fan el idioma de la factura. Guardarlos como texto evita problemas de parseo.'
)

doc.add_heading(u'4.2. Servicio de base de datos (ServicioBaseDatos)', level=2)

doc.add_paragraph(
    u'El servicio ServicioBaseDatos.cs maneja toda la comunicaci\u00f3n con SQLite. Usa el paquete '
    u'NuGet sqlite-net-pcl que funciona con async/await. Las operaciones que ofrece son:'
)

db_ops = [
    u'SaveInvoiceAsync: guarda o actualiza una factura. Si el Id es 0 hace INSERT, si no UPDATE.',
    u'GetAllInvoicesAsync: devuelve todas las facturas ordenadas por fecha (las m\u00e1s recientes primero).',
    u'GetInvoiceAsync: busca una factura por su Id para mostrar en la pantalla de detalle.',
    u'DeleteInvoiceAsync: elimina una factura de la BD.',
    u'FindDuplicateAsync: busca si ya existe una factura con el mismo nombre de archivo o con contenido similar (comparando los primeros 200 caracteres del texto). Esto se usa para evitar guardar duplicados.',
]
for op in db_ops:
    doc.add_paragraph(f'\u2022 {op}')

doc.add_paragraph('')

doc.add_paragraph(
    u'La ruta de la base de datos se puede cambiar desde los ajustes de la app. Por defecto '
    u'se guarda en una carpeta "DatosFacturas" junto al ejecutable, pero el usuario puede elegir '
    u'cualquier carpeta (por ejemplo el Escritorio). Si se cambia la ruta, la app detecta que '
    u'ha cambiado y reconecta a la nueva BD autom\u00e1ticamente.'
)


# ============== 5. REQUISITOS E INSTALACION ==============
doc.add_page_break()
doc.add_heading(u'5. Requisitos e instalaci\u00f3n', level=1)

doc.add_heading('5.1. Requisitos previos', level=2)

doc.add_paragraph(
    u'Para poder compilar y ejecutar el proyecto se necesita:'
)

requisitos = [
    ('Visual Studio 2022 (17.8 o superior)', u'Con la carga de trabajo ".NET Multi-platform App UI development" instalada.'),
    ('.NET 9 SDK', u'Se puede descargar desde https://dotnet.microsoft.com/download. El proyecto usa net9.0.'),
    ('Windows 10 (build 17763 o superior)', u'Para ejecutar la versi\u00f3n Windows de la app. Tambi\u00e9n se puede ejecutar en Android con un emulador.'),
    ('Cuenta de Azure con recurso Document Intelligence', u'Necesario para que el OCR funcione. Se puede usar el tier gratuito F0.'),
]

for titulo_req, desc_req in requisitos:
    p = doc.add_paragraph()
    run_bold = p.add_run(f'\u2022 {titulo_req}: ')
    run_bold.font.bold = True
    run_bold.font.size = Pt(10)
    run_normal = p.add_run(desc_req)
    run_normal.font.size = Pt(10)

doc.add_paragraph('')

doc.add_heading(u'5.2. C\u00f3mo ejecutar el proyecto', level=2)

doc.add_paragraph(
    u'Pasos para abrir y ejecutar el proyecto:'
)

pasos_instalacion = [
    u'Abrir el archivo FacturasOCR.sln (o FacturasOCR.csproj) con Visual Studio 2022.',
    u'Visual Studio deber\u00eda restaurar autom\u00e1ticamente los paquetes NuGet. Si no, hacer clic derecho en el proyecto \u2192 "Restaurar paquetes NuGet".',
    u'Seleccionar la plataforma de destino en la barra de herramientas. Para Windows: "Windows Machine". Para Android: seleccionar un emulador o dispositivo conectado.',
    u'Pulsar F5 o el bot\u00f3n de Play para compilar y ejecutar.',
]

for i, paso_inst in enumerate(pasos_instalacion):
    doc.add_paragraph(f'{i+1}. {paso_inst}')

doc.add_paragraph('')

doc.add_paragraph(
    u'Tambi\u00e9n se puede compilar desde la terminal con el comando:'
)

p_cmd = doc.add_paragraph()
p_cmd.paragraph_format.left_indent = Cm(1)
run = p_cmd.add_run('dotnet build -f net9.0-windows10.0.19041.0')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph('')

doc.add_paragraph(
    u'Los paquetes NuGet que usa el proyecto son:'
)

table_nuget = doc.add_table(rows=5, cols=2, style='Table Grid')
table_nuget.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table_nuget.rows[0].cells
hdr[0].text = 'Paquete NuGet'
hdr[1].text = u'Versi\u00f3n'
for cell in table_nuget.rows[0].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

nugets = [
    ('Azure.AI.FormRecognizer', '4.1.0'),
    ('sqlite-net-pcl', '1.9.172'),
    ('Microsoft.Identity.Client', '4.67.2'),
    ('Microsoft.Maui.Controls', '(incluido con .NET 9)'),
]
for i, (pkg, ver) in enumerate(nugets):
    cells = table_nuget.rows[i+1].cells
    cells[0].text = pkg
    cells[0].paragraphs[0].runs[0].font.name = 'Consolas'
    cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    cells[1].text = ver

aplicar_fuente_tabla(table_nuget)
doc.add_paragraph('')

doc.add_heading(u'5.3. Configuraci\u00f3n de Azure', level=2)

doc.add_paragraph(
    u'La app viene con un endpoint y API key por defecto que son los del recurso creado para esta '
    u'pr\u00e1ctica. Si se quiere usar otro recurso de Azure, se pueden cambiar desde la pesta\u00f1a '
    u'de Ajustes dentro de la app, o directamente editando el archivo '
    u'Servicios/ServicioAnalisis.cs.'
)

doc.add_paragraph(
    u'Para crear un recurso nuevo de Document Intelligence en Azure:'
)

pasos_azure = [
    u'Entrar en portal.azure.com con la cuenta de estudiante.',
    u'Buscar "Document Intelligence" en la barra de b\u00fasqueda.',
    u'Crear un nuevo recurso con el tier F0 (gratuito).',
    u'Una vez creado, ir a "Keys and Endpoint" y copiar el Endpoint y la Key 1.',
    u'Pegar esos valores en la pantalla de Ajustes de la app.',
]

for i, paso_az in enumerate(pasos_azure):
    doc.add_paragraph(f'{i+1}. {paso_az}')


# ============== 6. FUNCIONALIDADES ==============
doc.add_page_break()
doc.add_heading('6. Funcionalidades principales', level=1)

doc.add_heading('6.1. Procesamiento de facturas', level=2)
doc.add_paragraph(
    u'La app puede procesar tres tipos de archivos:'
)
doc.add_paragraph(u'\u2022 Im\u00e1genes (JPG, PNG, BMP, TIFF): se mandan directamente a Azure Document Intelligence.')
doc.add_paragraph(u'\u2022 PDFs: se mandan a Azure igual que las im\u00e1genes, el servicio sabe leerlos.')
doc.add_paragraph(u'\u2022 XMLs de factura electr\u00f3nica (FacturaE, XSIG): se parsean localmente sin necesidad de usar Azure, leyendo directamente las etiquetas XML.')

doc.add_paragraph('')
doc.add_paragraph(
    u'Para las im\u00e1genes y PDFs, despu\u00e9s de recibir el AnalyzeResult de Azure (igual que en '
    u'el ejemplo de clase), la app primero intenta sacar los campos de Documents[].Fields '
    u'(VendorName, InvoiceId, etc.), que es lo que en los apuntes se llama "Entities/Fields" '
    u'(ShowEntitiesSummary). Si Azure no los detecta (porque la factura es rara o no est\u00e1 bien '
    u'escaneada), la app hace una b\u00fasqueda manual en los pares clave-valor (Key-Value Pairs) '
    u'y en las tablas como fallback, usando la misma estructura que las funciones '
    u'ShowKeyValuePairsSummary y ShowTablesSummary del ejemplo.'
)

doc.add_paragraph('')
doc.add_paragraph(
    u'Para entender mejor c\u00f3mo funciona, este es el orden exacto de b\u00fasqueda de cada campo '
    u'(por ejemplo para sacar el total de la factura):'
)

busqueda_pasos = [
    u'Primero se busca en Documents[0].Fields["InvoiceTotal"]. Si Azure lo ha detectado como campo del modelo prebuilt-invoice, devuelve el valor directamente con su tipo (Currency) y su nivel de confianza.',
    u'Si no lo encuentra ah\u00ed (porque la factura est\u00e1 en un formato raro o mal escaneada), se busca en los pares clave-valor (KeyValuePairs). Se comparan las claves con palabras como "total", "importe total", "amount due", "grand total", etc.',
    u'Si tampoco est\u00e1 en los pares clave-valor, se busca en las tablas (Tables). Se recorre cada celda buscando la palabra "total" y si la encuentra, coge el valor de la \u00faltima columna de esa misma fila (que suele ser el importe).',
    u'Si no se encuentra en ninguna parte, el campo se muestra como "(no detectado)" en la ficha visual.',
]
for i, paso_b in enumerate(busqueda_pasos):
    doc.add_paragraph(f'{i+1}. {paso_b}')

doc.add_paragraph('')
doc.add_paragraph(
    u'Este sistema de fallback en cascada es lo que hace que la app funcione razonablemente bien '
    u'con todo tipo de facturas, no solo con las que Azure reconoce perfectamente.'
)

doc.add_heading('6.2. Procesamiento por lotes', level=2)
doc.add_paragraph(
    u'Adem\u00e1s del procesamiento individual, la app tiene un bot\u00f3n "Procesar Lote" que permite '
    u'seleccionar varios archivos a la vez. Los procesa uno por uno de forma secuencial y al final '
    u'muestra un resumen con cu\u00e1ntos se procesaron bien y cu\u00e1ntos dieron error. Tambi\u00e9n '
    u'detecta duplicados autom\u00e1ticamente (compara el nombre de archivo y el contenido) para no '
    u'guardar la misma factura dos veces.'
)

doc.add_heading('6.3. Historial y detalle', level=2)
doc.add_paragraph(
    u'En la pesta\u00f1a "Archivo" se muestra la lista de todas las facturas procesadas, ordenadas por '
    u'fecha. Se puede tocar cualquiera para ver su detalle completo: la imagen original, los campos '
    u'extra\u00eddos (proveedor, n\u00famero, fecha, total) y el texto completo del OCR. Desde el '
    u'detalle se puede compartir la informaci\u00f3n o eliminar la factura de la base de datos.'
)

doc.add_heading('6.4. Sistema multiidioma', level=2)
doc.add_paragraph(
    u'La app soporta 9 idiomas que se pueden cambiar en tiempo real desde la pesta\u00f1a de Ajustes. '
    u'Al cambiar el idioma, todos los textos de toda la app se actualizan instant\u00e1neamente '
    u'sin necesidad de reiniciar. Los idiomas disponibles son:'
)

idiomas = [
    u'Castellano (idioma por defecto)',
    'English',
    u'Catal\u00e0',
    'Euskara',
    'Galego',
    u'Aragon\u00e9s',
    u'Aran\u00e9s (occitano)',
    'Asturianu',
    u'Extreme\u00f1u',
]

for idioma in idiomas:
    doc.add_paragraph(f'\u2022 {idioma}')

doc.add_paragraph('')
doc.add_paragraph(
    u'El sistema de traducciones funciona con un servicio est\u00e1tico (ServicioIdiomas) que '
    u'tiene un diccionario con todas las claves de texto y sus traducciones. Cuando se cambia el idioma '
    u'se lanza un evento que todas las p\u00e1ginas escuchan para actualizar sus textos.'
)

doc.add_heading('6.5. Tema claro / oscuro', level=2)
doc.add_paragraph(
    u'Desde la pesta\u00f1a de Ajustes se puede cambiar el tema visual de la aplicaci\u00f3n entre '
    u'tres opciones: Claro, Oscuro o Seguir el sistema (usa el tema que tenga configurado '
    u'el dispositivo). El tema se guarda en las preferencias de la app y se aplica autom\u00e1ticamente '
    u'al arrancar, sin necesidad de volver a configurarlo cada vez.'
)


# ============== 7. CAPTURAS ==============
doc.add_page_break()
doc.add_heading('7. Capturas de pantalla', level=1)

doc.add_paragraph(
    u'A continuaci\u00f3n se muestran capturas de la aplicaci\u00f3n funcionando:'
)

capturas_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'capturas')

capturas = [
    ('Captura de pantalla 2026-03-03 110314.png', u'Pantalla principal antes de procesar ninguna factura'),
    ('Captura de pantalla 2026-03-03 110545.png', u'Procesando una factura con Azure (spinner de carga)'),
    ('Captura de pantalla 2026-03-03 110604.png', u'Resultado del OCR: ficha visual con los campos extra\u00eddos'),
    ('Captura de pantalla 2026-03-03 110654.png', u'Detalle del texto completo devuelto por Azure'),
    ('Captura de pantalla 2026-03-03 110734.png', u'Historial de facturas guardadas'),
    ('Captura de pantalla 2026-03-03 110812.png', u'Detalle de una factura desde el historial'),
    ('Captura de pantalla 2026-03-03 110844.png', u'Pantalla de ajustes: Azure y apariencia'),
    ('Captura de pantalla 2026-03-03 110903.png', u'Pantalla de ajustes: almacenamiento, backup e idioma'),
    ('Captura de pantalla 2026-03-03 111008.png', u'Ejemplo de la app en euskara'),
]

for i, (archivo, caption) in enumerate(capturas):
    img_path = os.path.join(capturas_dir, archivo)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.5))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f'[ CAPTURA NO ENCONTRADA: {archivo} ]')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(180, 180, 180)
        run.font.italic = True

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(16)
    run = cap.add_run(f'Fig. {i+2}: {caption}')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)


# ============== 8. CONCLUSIONES ==============
doc.add_page_break()
doc.add_heading('8. Conclusiones', level=1)

doc.add_paragraph(
    u'La aplicaci\u00f3n cumple con los requisitos del Proyecto 2 de la tarea: permite hacer fotos '
    u'de facturas con la c\u00e1mara o elegir archivos, procesarlos con Azure Document Intelligence '
    u'y almacenar los datos extra\u00eddos como texto plano. Los materiales de clase (sobre todo el '
    u'c\u00f3digo de ejemplo del FormRecognizer, el documento de creaci\u00f3n del recurso y el glosario '
    u'de t\u00e9rminos) me han sido muy \u00fatiles como punto de partida para entender c\u00f3mo funciona '
    u'todo el ecosistema de Azure Cognitive Services.'
)

doc.add_paragraph(
    u'Lo que m\u00e1s me ha costado ha sido la parte de extracci\u00f3n de campos, porque el ejemplo '
    u'de clase usa "prebuilt-document" (gen\u00e9rico) y yo necesitaba "prebuilt-invoice" (espec\u00edfico '
    u'de facturas), que devuelve los campos de forma diferente. Adem\u00e1s cada factura tiene un '
    u'formato diferente y a veces Azure no detecta todos los campos, as\u00ed que implement\u00e9 un '
    u'fallback que busca manualmente en los Key-Value Pairs y en las Tables, que es b\u00e1sicamente '
    u'lo mismo que hacen las funciones Show del ejemplo pero adaptado a mi caso.'
)

doc.add_paragraph(
    u'Tambi\u00e9n ha sido interesante implementar el soporte multiidioma, que aunque no era un '
    u'requisito obligatorio, me pareci\u00f3 una buena forma de practicar con eventos y '
    u'actualizaci\u00f3n din\u00e1mica de la interfaz.'
)

doc.add_paragraph(
    u'Como posibles mejoras futuras se podr\u00eda a\u00f1adir exportaci\u00f3n a Excel de las '
    u'facturas guardadas, o integraci\u00f3n con alg\u00fan servicio de almacenamiento en la nube '
    u'como OneDrive para sincronizar las facturas entre dispositivos.'
)


# ---------------------------------------------------
# Guardar
# ---------------------------------------------------
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Tarea 7 - Documento - autor.docx')
doc.save(output_path)
print(f'\nDocumento guardado en:\n{output_path}')

# Limpiar diagrama temporal
if diagrama_path and os.path.exists(diagrama_path):
    os.remove(diagrama_path)
    print('Diagrama temporal eliminado')

print('\nLISTO! Abre el .docx en Word, inserta las capturas y exporta a PDF.')
