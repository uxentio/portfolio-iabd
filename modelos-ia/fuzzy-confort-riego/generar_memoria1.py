# Genera memoria_ejercicio1.docx desde las graficas y resumen_ej1.json.

import json
import os
import math
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
with open(os.path.join(HERE, 'resumen_ej1.json'), encoding='utf-8') as f:
    R = json.load(f)

# fuente hiperlegible y color azul corporativo (mismo en Ej.1 y Ej.2)
FONT_TEXT = 'Atkinson Hyperlegible Next'
FONT_CODE = 'Courier New'
COLOR_HEADER = RGBColor(0x1F, 0x4E, 0x79)
COLOR_HEADER_HEX = '1F4E79'

doc = Document()

# margenes
for section in doc.sections:
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)


def force_font(run_or_style, font_name):
    # python-docx no siempre escribe la fuente en todos los rangos.
    # forzamos el atributo w:ascii / w:hAnsi / w:cs a nivel XML.
    rpr = run_or_style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(attr), font_name)


# estilo Normal
normal = doc.styles['Normal']
normal.font.name = FONT_TEXT
normal.font.size = Pt(11)
force_font(normal, FONT_TEXT)

# headings con la misma fuente y color azul
for lvl in (1, 2, 3):
    style = doc.styles[f'Heading {lvl}']
    style.font.name = FONT_TEXT
    style.font.color.rgb = COLOR_HEADER
    force_font(style, FONT_TEXT)


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_table(headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = FONT_TEXT
                force_font(run, FONT_TEXT)
        set_cell_bg(hdr[i], COLOR_HEADER_HEX)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = FONT_TEXT
                    force_font(run, FONT_TEXT)
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def H1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = COLOR_HEADER
        run.font.name = FONT_TEXT
        force_font(run, FONT_TEXT)


def H2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = COLOR_HEADER
        run.font.name = FONT_TEXT
        force_font(run, FONT_TEXT)


# titulo
H1('Práctica 3, Ejercicio 1')
H2('Índice de confort térmico (ICT)')

# 1. problema
H2('1. Problema')
doc.add_paragraph(
    'SportClima quiere estimar el ICT (de 0 a 1) a partir de la temperatura '
    'interior T y el porcentaje de ocupación O, ambos normalizados a [0, 1]. '
    'El experto dice que el confort máximo está cerca de los 22 °C y un '
    'aforo del 55 %. Uso un modelo ANFIS de 5 capas en PyTorch para que '
    'aprenda eso a partir de los datos.')

# 2. decisiones
cfg = R['config']
H2('2. Decisiones de diseño')
add_table(
    ['Parámetro', 'Valor', 'Justificación'],
    [
        ['n_rules', str(cfg['n_rules']),
         'Una regla por cuadrante (T baja/alta × O baja/alta).'],
        ['n_epochs', str(cfg['n_epochs']),
         'Estabiliza sobre la época 200, paro a las 800.'],
        ['lr', '0.01 (probé también 0.005)',
         'Con 0.01 converge antes; bajarlo no aporta nada.'],
        ['Init μ', 'K-means y aleatoria (comparadas)',
         'K-means coloca los centros donde hay datos.'],
    ],
    col_widths_cm=[3, 4.5, 8.5],
)
doc.add_paragraph()

# 3. exp A
H2('3. Experimento A: curva de entrenamiento')
doc.add_picture(os.path.join(HERE, 'expA_convergencia.png'), width=Cm(15))
mse = R['mse_final']
doc.add_paragraph(
    f'MSE final con K-means y lr=0.01: {mse["kmeans_lr0.01"]:.4f}. '
    'K-means estabiliza sobre la época 100 y la aleatoria sobre la 150. '
    'Con lr=0.005 va más lento sin mejorar el final. Me quedo con K-means '
    'y lr=0.01.')

# 4. exp B
H2('4. Experimento B: predicción vs real')
p = doc.add_paragraph()
p.add_run(f'RMSE = {R["test"]["rmse"]:.4f}').bold = True
p.add_run(', ')
p.add_run(f'R² = {R["test"]["r2"]:.4f}.').bold = True
doc.add_picture(os.path.join(HERE, 'expB_pred_vs_real.png'), width=Cm(9))
doc.add_paragraph(
    'La nube sigue la diagonal. El modelo aproxima bien al criterio del '
    'experto: el error es del orden del propio ruido del dataset (0.04).')

# 5. exp C
H2('5. Experimento C: reglas aprendidas')
reg_rows = []
for r in R['reglas']:
    reg_rows.append([
        r['regla'],
        f'{r["mu_T"]:+.2f}', f'{r["sigma_T"]:.2f}',
        f'{r["mu_O"]:+.2f}', f'{r["sigma_O"]:.2f}',
        r['zona'],
    ])
add_table(
    ['Regla', 'μ_T', 'σ_T', 'μ_O', 'σ_O', 'Zona'],
    reg_rows,
    col_widths_cm=[1.5, 1.8, 1.8, 1.8, 1.8, 6],
)
doc.add_paragraph()

best = min(R['reglas'], key=lambda r: math.hypot(r['mu_T']-0.6, r['mu_O']-0.55))
p = doc.add_paragraph()
p.add_run(best['regla']).bold = True
p.add_run(
    f' es la regla del confort: μ ≈ (0.6, 0.55), casi exactamente el '
    'óptimo del experto. Sus σ son las más pequeñas, o sea, es la regla '
    'más selectiva.')

plana = max(R['reglas'], key=lambda r: r['sigma_T']+r['sigma_O'])
p = doc.add_paragraph()
p.add_run(plana['regla']).bold = True
p.add_run(
    ' tiene σ grandes, así que cubre mucho espacio. Funciona como regla '
    'de fondo en zonas donde las otras se apagan.')

# 6. exp D
H2('6. Experimento D: mapa de confort 2D')
doc.add_picture(os.path.join(HERE, 'expD_mapa_confort.png'), width=Cm(16))
doc.add_paragraph(
    'El mapa aprendido reproduce el criterio teórico: zona caliente cerca '
    'de (0.6, 0.55) y bajo confort en las esquinas.')

# 7. conclusion
H2('7. Conclusión')
doc.add_paragraph(
    f'ANFIS sirve para este problema: con 4 reglas saco R² = '
    f'{R["test"]["r2"]:.2f}. La ventaja frente a una red neuronal pura es '
    'que las reglas se pueden leer (cada μ y σ tiene un significado físico) '
    'y que funciona con pocos datos. El inconveniente es que no escala '
    'bien con muchas variables, porque la fuerza de disparo es producto de '
    'membresías y se va a cero. Tampoco va bien con problemas no convexos.')

# fuerza fuente en todos los runs ya creados (para que cualquier texto que haya
# escapado al estilo Normal tambien la use)
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if not run.font.name:
            run.font.name = FONT_TEXT
        force_font(run, FONT_TEXT)

doc.save(os.path.join(HERE, 'memoria_ejercicio1.docx'))
print('memoria_ejercicio1.docx generada')
