# Genera memoria_ejercicio2.docx desde las graficas y resumen_ej2.json.

import json
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
with open(os.path.join(HERE, 'resumen_ej2.json'), encoding='utf-8') as f:
    R = json.load(f)

# fuente hiperlegible y color azul corporativo (mismo en Ej.1 y Ej.2)
FONT_TEXT = 'Atkinson Hyperlegible Next'
FONT_CODE = 'Courier New'
COLOR_HEADER = RGBColor(0x1F, 0x4E, 0x79)
COLOR_HEADER_HEX = '1F4E79'

doc = Document()

for section in doc.sections:
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)


def force_font(run_or_style, font_name):
    rpr = run_or_style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(attr), font_name)


# estilo Normal con fuente hiperlegible
normal = doc.styles['Normal']
normal.font.name = FONT_TEXT
normal.font.size = Pt(11)
force_font(normal, FONT_TEXT)

for lvl in (1, 2, 3):
    style = doc.styles[f'Heading {lvl}']
    style.font.name = FONT_TEXT
    style.font.color.rgb = COLOR_HEADER
    force_font(style, FONT_TEXT)


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_table(headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = FONT_TEXT
                force_font(run, FONT_TEXT)
        set_cell_bg(hdr[i], COLOR_HEADER_HEX)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = FONT_TEXT
                    force_font(run, FONT_TEXT)
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def H1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = COLOR_HEADER
        run.font.name = FONT_TEXT
        force_font(run, FONT_TEXT)


def H2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = COLOR_HEADER
        run.font.name = FONT_TEXT
        force_font(run, FONT_TEXT)


def H3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = COLOR_HEADER
        run.font.name = FONT_TEXT
        force_font(run, FONT_TEXT)


def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_CODE
    run.font.size = Pt(9)
    force_font(run, FONT_CODE)


# titulo
H1('Práctica 3, Ejercicio 2')
H2('Sistema de riego adaptativo')

# 1. problema
H2('1. Problema')
doc.add_paragraph(
    'AgroSmart Levante quiere recomendar la dosis de riego (en mm/día) a '
    'partir de tres sensores: humedad del suelo H, temperatura máxima T y '
    'precipitación P. Las tres están normalizadas a [0, 1] y la salida '
    'también. Uso ANFIS con 3 entradas y la versión robusta del modelo '
    '(con clamp en σ).')

# 2. funcion objetivo
H2('2. Función objetivo')
doc.add_paragraph('Los valores con «?» del enunciado los relleno así:')
code_block(
    'efecto_lluvia  = exp(-((P - 1.0)^2 / 0.30))\n'
    'efecto_humedad = exp(-((H - 1.0)^2 / 0.30))\n'
    'efecto_calor   = exp(-((T - 1.0)^2 / 0.20))'
)
doc.add_paragraph(
    'Centro las tres en 1.0 porque cada efecto es un indicador de '
    '«extremo activo» (lluvia fuerte, suelo saturado, calor). Multiplicar '
    'por (1 − efecto) reproduce los vetos del agrónomo: si llueve o el '
    'suelo está saturado, la dosis se anula. σ_T más estrecha (0.20) '
    'porque el calor solo cuenta cuando T > 0.6.')

# decisiones
cfg = R['config']
add_table(
    ['Parámetro', 'Valor', 'Justificación'],
    [
        ['n_rules', str(cfg['n_rules_elegido']),
         'Mejor RMSE en el experimento A.'],
        ['n_epochs', str(cfg['n_epochs']),
         'Estabiliza por la época 200.'],
        ['lr', str(cfg['lr']),
         'Más bajo que en el Ej.1 (3 factores en el producto).'],
        ['weight_decay', str(cfg['weight_decay']),
         'Regularización L2 obligatoria.'],
        ['Init μ', 'K-means', 'Obligatorio con 3 entradas.'],
        ['Forward', 'torch.clamp(σ, min=1e-3)',
         'Modificación obligatoria, evita que las gaussianas colapsen.'],
    ],
    col_widths_cm=[3, 4, 9],
)
doc.add_paragraph()

# 3. ANFISLayerRobust
H2('3. Por qué hace falta ANFISLayerRobust')
doc.add_paragraph(
    'Con 3 entradas, w = μ_A · μ_B · μ_C. Si σ se acerca a 0, la gaussiana '
    'se hace un pico estrecho y w se va a 0 para casi todas las muestras. '
    'Entonces el denominador de la normalización también es 0 y los '
    'gradientes salen NaN: el modelo deja de entrenar. El clamp obliga a '
    'σ ≥ 1e-3, así que la gaussiana siempre tiene un soporte mínimo.')

# 4. exp A
H2('4. Experimento A: distintos n_rules')
A = R['expA']
expA_rows = []
for nr_str in ['3', '5', '8']:
    info = A[nr_str]
    n_params = int(nr_str) * (3 * 2 + 4)
    expA_rows.append([nr_str, str(n_params),
                      f'{info["rmse"]:.4f}', f'{info["r2"]:.4f}'])
add_table(
    ['n_rules', 'Parámetros', 'RMSE (test)', 'R² (test)'],
    expA_rows,
    col_widths_cm=[2, 4, 4, 4],
)
doc.add_paragraph()
doc.add_picture(os.path.join(HERE, 'expA_n_rules.png'), width=Cm(15))
doc.add_paragraph(
    f'Me quedo con n_rules = 8 (RMSE = {A["8"]["rmse"]:.4f}). Con 3 se '
    'queda corto y con 5 va competitivo, pero 8 mejora un poco más sin '
    'sobreajustarse.')

# 5. exp B
H2('5. Experimento B: aleatoria vs K-means')
B = R['expB']
doc.add_picture(os.path.join(HERE, 'expB_init.png'), width=Cm(15))
add_table(
    ['Init', 'RMSE test', 'Épocas hasta el régimen final'],
    [
        ['Aleatoria', f'{B["rmse_aleatoria"]:.4f}', str(B["epocas_aleatoria"])],
        ['K-means', f'{B["rmse_kmeans"]:.4f}', str(B["epocas_kmeans"])],
    ],
    col_widths_cm=[4, 3, 8],
)
doc.add_paragraph()
doc.add_paragraph(
    f'K-means ahorra unas {B["epocas_ahorradas"]} épocas '
    f'(~{B["ahorro_pct"]:.0f} %) y además da mejor RMSE de test. Con la '
    'aleatoria, los centros caen lejos de los datos y muchas reglas '
    'arrancan con fuerza casi nula.')

# 6. exp C
H2('6. Experimento C: importancia de variables')
C = R['expC']
sm = C['sigma_media_por_variable']
add_table(
    ['Variable', 'σ media (8 reglas)', 'Interpretación'],
    [
        ['H (humedad)', f'{sm["H"]:.3f}', '2.ª más selectiva.'],
        ['T (temperatura)', f'{sm["T"]:.3f}', 'Menos selectiva.'],
        ['P (precipitación)', f'{sm["P"]:.3f}', 'La más selectiva.'],
    ],
    col_widths_cm=[4, 4, 7],
)
doc.add_paragraph()
doc.add_paragraph(
    'P sale como la más determinante (σ media más pequeña). Coincide solo '
    'en parte con el agrónomo: él dice que H es la más importante en '
    'situaciones intermedias, pero P es la regla de veto más fuerte '
    '(«si P > 0.6, no riego»). Cuando llueve fuerte, las demás dejan de '
    'importar.')

# 7. exp D
H2('7. Experimento D: vértices del cubo')
D = R['expD']
vert_rows = []
for row in D['tabla_vertices']:
    vert_rows.append([
        f'{row["H"]:.0f}', f'{row["T"]:.0f}', f'{row["P"]:.0f}',
        f'{row["pred_anfis"]:.3f}',
        f'{row["experto"]:.3f}',
        f'{row["error_abs"]:.3f}',
    ])
add_table(
    ['H', 'T', 'P', 'Pred ANFIS', 'Experto', '|err|'],
    vert_rows,
    col_widths_cm=[1, 1, 1, 3, 3, 2],
)
doc.add_paragraph()
doc.add_paragraph(
    f'Error medio {D["error_medio"]:.3f}, máximo {D["error_max"]:.3f} en '
    '(H=0, T=1, P=1): ahí chocan dos reglas (suelo seco + calor pide '
    'regar, pero la lluvia dice que no). El experto resuelve por '
    'prioridad y ANFIS promedia, así que predice 0.139 en vez de 0. En '
    'el resto, el error está por debajo del ruido del sensor.')

# 8. exp E
H2('8. Experimento E: sistema explicable')
doc.add_paragraph(
    'Adapto la función explicar_decision() del documento de referencia '
    'a 3 entradas. Devuelve la dosis y la activación de cada regla.')

ca = R['expE']['caso_a']
p = doc.add_paragraph()
p.add_run('(a) Tarde de julio seca').bold = True
p.add_run(
    f' (H={ca["H"]:.2f}, T={ca["T"]:.2f}, P={ca["P"]:.2f}): dosis = '
    f'{ca["dosis_norm"]:.3f} ({ca["dosis_mm_dia"]:.2f} mm/día), ')
p.add_run(ca['nivel']).bold = True
p.add_run('. Coincide con lo que diría el agrónomo (regar mucho).')

cb = R['expE']['caso_b']
p = doc.add_paragraph()
p.add_run('(b) Mañana lluviosa de noviembre').bold = True
p.add_run(
    f' (H={cb["H"]:.2f}, T={cb["T"]:.2f}, P={cb["P"]:.2f}): dosis = '
    f'{cb["dosis_norm"]:.3f}, ')
p.add_run(cb['nivel']).bold = True
p.add_run('. Doble veto activo (suelo saturado y lluvia), predice cero.')

# 9. conclusion
H2('9. Conclusión')
fin = R['final']
doc.add_paragraph(
    f'ANFIS frente a una red neuronal pura, en agricultura yo elegiría '
    f'ANFIS. Con 8 reglas saco R² = {fin["r2_test"]:.2f} y, sobre todo, '
    'puedo enseñar al agrónomo qué reglas se han disparado en cada '
    'recomendación. Una MLP sería caja negra. Tampoco necesito muchos '
    'datos. Lo cambiaría por una red neuronal si tuviera 8 o 10 entradas '
    'o el problema fuera muy no convexo (por ejemplo, clasificar plagas '
    'en imágenes), porque ahí ANFIS escala mal.')

# pasada final: aplico la fuente hiperlegible a todo run que haya escapado
# (excepto los bloques de codigo, que ya tienen Courier New)
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.font.name == FONT_CODE:
            continue
        if not run.font.name:
            run.font.name = FONT_TEXT
        force_font(run, FONT_TEXT)

doc.save(os.path.join(HERE, 'memoria_ejercicio2.docx'))
print('memoria_ejercicio2.docx generada')
